"""
Creates Framework_Semanticas_Summary_IEEE_Final.docx from the corrected markdown source.
IEEE Wireless Communications journal article format.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH = "/home/runner/work/Papers/Papers/P9/Framework_Semanticas_Summary_IEEE.md"
OUTPUT_PATH = "/home/runner/work/Papers/Papers/P9/Framework_Semanticas_Summary_IEEE_Final.docx"

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)


# ── helpers ───────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=10, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Ensure font applies to East Asian chars too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rPr.insert(0, rFonts)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=14, bold=True)


def add_authors(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=10, italic=True)


def add_heading1(text):
    """Section heading – Roman numerals, 11pt bold centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=11, bold=True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    return p


def add_heading2(text):
    """Subsection heading – letter prefix, 10pt bold."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10, bold=True)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    return p


def add_heading3(text):
    """Sub-subsection heading – 10pt bold italic."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=10, bold=True, italic=True)
    p.paragraph_format.space_before = Pt(3)
    return p


def add_abstract(text):
    """Abstract paragraph – 9pt italic, justified, prefixed 'Abstract—'."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run("Abstract—")
    set_font(run_label, size=9, bold=True, italic=True)
    run_body = p.add_run(text)
    set_font(run_body, size=9, italic=True)


def add_keywords(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run("Index Terms—")
    set_font(run_label, size=9, bold=True, italic=True)
    run_body = p.add_run(text)
    set_font(run_body, size=9, italic=True)


def add_figure_caption(text):
    """Figure caption – 9pt italic, slightly indented."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_font(run, size=9, italic=True)


def add_table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=9, bold=True)


def add_algo_line(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_reference(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent    = Inches(0.35)
    fmt.first_line_indent = Inches(-0.35)
    _add_inline_runs(p, text, base_size=9)
    return p


def _add_inline_runs(p, text, base_size=10):
    """Parse **bold**, *italic*, inline $math$ and add runs to paragraph p."""
    # Tokenize: bold (**...**), italic (*...*), inline math ($...$)
    pattern = re.compile(r'(\*\*.*?\*\*|\*[^*]+?\*|\$[^$]+?\$)')
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            run = p.add_run(text[pos:m.start()])
            set_font(run, size=base_size)
        token = m.group(0)
        if token.startswith('**'):
            run = p.add_run(token[2:-2])
            set_font(run, size=base_size, bold=True)
        elif token.startswith('*'):
            run = p.add_run(token[1:-1])
            set_font(run, size=base_size, italic=True)
        elif token.startswith('$'):
            # Keep math as plain text
            math_text = token[1:-1]
            run = p.add_run(math_text)
            set_font(run, size=base_size)
        pos = m.end()
    # remaining plain text
    if pos < len(text):
        run = p.add_run(text[pos:])
        set_font(run, size=base_size)


def add_body_paragraph(text, justified=True):
    """Body text with inline formatting."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    if justified:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_inline_runs(p, text, base_size=10)
    return p


def add_math_block(text):
    """Standalone math block as body paragraph."""
    # Strip $$ delimiters
    inner = text.strip().strip('$').strip()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(inner)
    set_font(run, size=10, italic=True)


def render_md_table(table_lines):
    """Parse markdown table lines and add a Word table."""
    rows = []
    for line in table_lines:
        if re.match(r'\|[-: |]+\|', line):
            continue  # separator row
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    headers = rows[0]
    data    = rows[1:]
    ncols   = len(headers)
    nrows   = 1 + len(data)
    t = doc.add_table(rows=nrows, cols=ncols)
    t.style = "Table Grid"
    for ci, h in enumerate(headers):
        cell = t.cell(0, ci)
        cell.text = ""
        rn = cell.paragraphs[0].add_run(h)
        set_font(rn, size=9, bold=True)
    for ri, row in enumerate(data):
        for ci in range(ncols):
            cell = t.cell(ri + 1, ci)
            cell.text = ""
            txt = row[ci] if ci < len(row) else ""
            _add_inline_runs(cell.paragraphs[0], txt, base_size=9)
    doc.add_paragraph()  # spacer


# ── read markdown ─────────────────────────────────────────────────────────────

with open(MD_PATH, encoding="utf-8") as f:
    raw_lines = f.readlines()

lines = [ln.rstrip('\n') for ln in raw_lines]

# ── document header ───────────────────────────────────────────────────────────

add_title("A Multi-Dimensional Semantic Metric Standardization Framework for "
          "Evaluating AI-Native Systems in 6G Networks")
doc.add_paragraph()
add_authors("Author Name¹, Author Name², Author Name³")
doc.add_paragraph()

# ── state machine parser ──────────────────────────────────────────────────────

i = 0
in_abstract   = False
abstract_done = False
table_buffer  = []
skip_title    = True   # skip the H1 title line (already added above)

while i < len(lines):
    line = lines[i]

    # Skip the very first H1 title
    if skip_title and line.startswith("# "):
        skip_title = False
        i += 1
        continue

    # Horizontal rule
    if re.match(r'^---+$', line.strip()):
        i += 1
        continue

    # ── TABLE collection ──────────────────────────────────────────────────────
    if line.startswith('|'):
        table_buffer.append(line)
        i += 1
        while i < len(lines) and lines[i].startswith('|'):
            table_buffer.append(lines[i])
            i += 1
        render_md_table(table_buffer)
        table_buffer = []
        continue

    # ── MATH BLOCK $$...$$ ────────────────────────────────────────────────────
    if line.strip().startswith('$$'):
        math_lines = []
        # Collect until closing $$
        rest = line.strip()[2:]
        if rest.endswith('$$') and len(rest) > 2:
            # single-line $$..$$
            add_math_block(rest[:-2])
            i += 1
            continue
        math_lines.append(rest)
        i += 1
        while i < len(lines):
            l2 = lines[i].strip()
            if l2 == '$$' or l2.endswith('$$'):
                math_lines.append(l2.rstrip('$').rstrip())
                i += 1
                break
            math_lines.append(l2)
            i += 1
        add_math_block('\n'.join(math_lines))
        continue

    # ── SECTION HEADINGS ──────────────────────────────────────────────────────
    if line.startswith('#### '):
        add_heading3(line[5:])
        i += 1
        continue

    if line.startswith('### '):
        add_heading2(line[4:])
        i += 1
        continue

    if line.startswith('## '):
        heading_text = line[3:].strip()
        # ABSTRACT handled specially
        if heading_text.upper() == 'ABSTRACT':
            in_abstract = True
            i += 1
            # Collect abstract body lines
            abstract_lines = []
            while i < len(lines):
                l2 = lines[i]
                if l2.startswith('##') or l2.startswith('**Index Terms'):
                    break
                if l2.strip():
                    abstract_lines.append(l2.strip())
                i += 1
            add_abstract(' '.join(abstract_lines))
            in_abstract = False
            abstract_done = True
            continue
        add_heading1(heading_text)
        i += 1
        continue

    # ── INDEX TERMS line ──────────────────────────────────────────────────────
    if line.startswith('**Index Terms'):
        # Extract content
        text = re.sub(r'^\*\*Index Terms\*\*—?', '', line).strip()
        text = text.strip('*').strip()
        add_keywords(text)
        i += 1
        continue

    # ── FIGURE CAPTIONS ───────────────────────────────────────────────────────
    if re.match(r'^\*\*Fig\.', line):
        # Strip ** markers
        text = re.sub(r'\*\*', '', line).strip()
        add_figure_caption(text)
        i += 1
        continue

    # ── TABLE CAPTION ─────────────────────────────────────────────────────────
    if re.match(r'^\*\*(TABLE|Table)\s', line):
        text = re.sub(r'^\*\*', '', line).rstrip('*').strip()
        add_table_caption(text)
        i += 1
        continue

    # ── ALGORITHM ─────────────────────────────────────────────────────────────
    if re.match(r'^\*\*Algorithm', line):
        text = re.sub(r'\*\*', '', line).strip()
        add_algo_line(text)
        i += 1
        continue

    # ── REFERENCE ENTRIES [N] ─────────────────────────────────────────────────
    if re.match(r'^\[\d+\]', line):
        add_reference(line)
        i += 1
        continue

    # ── EMPTY LINE ────────────────────────────────────────────────────────────
    if not line.strip():
        i += 1
        continue

    # ── BODY PARAGRAPH ────────────────────────────────────────────────────────
    add_body_paragraph(line)
    i += 1


# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
